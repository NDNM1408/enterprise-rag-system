"""Vendored from ragflow/deepdoc/parser/docx_parser.py.

Changes vs upstream:
  * dropped ``from rag.nlp import rag_tokenizer`` — the only callers were
    inside ``__compose_table_content`` for column type heuristics; replaced
    with a simple whitespace-based word counter (the heuristic is approximate
    either way).
  * dropped ``from rag.utils.lazy_image import LazyImage`` — image bytes are
    returned directly; no lazy wrapper.
"""
from __future__ import annotations

import logging
import re
from collections import Counter
from io import BytesIO

import pandas as pd
from docx import Document
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)


def _simple_tokenize(text: str) -> list[str]:
    """Drop-in replacement for ``rag_tokenizer.tokenize(...).split()``.

    The original tokenizer is jieba-based and tags POS; here we only need a
    word count >3 vs <12 to bucket cells. Whitespace + punctuation split is
    good enough for the purposes of the heuristic below.
    """
    return [t for t in re.split(r"[\s,，;；:：/\\\\|]+", text) if len(t) > 1]


class RAGFlowDocxParser:
    def get_picture(self, document, paragraph):
        imgs = paragraph._element.xpath(".//pic:pic")
        if not imgs:
            return None
        image_blobs = []
        for img in imgs:
            embed = img.xpath(".//a:blip/@r:embed")
            if not embed:
                continue
            embed = embed[0]
            image_blob = None
            try:
                related_part = document.part.related_parts[embed]
            except Exception as e:
                logging.warning(
                    "Skipping image due to unexpected error getting related_part: %s", e
                )
                continue
            try:
                image = related_part.image
                if image is not None:
                    image_blob = image.blob
            except (
                UnrecognizedImageError,
                UnexpectedEndOfFileError,
                InvalidImageStreamError,
                UnicodeDecodeError,
            ) as e:
                logging.info("Damaged image, attempting blob fallback: %s", e)
            except Exception as e:
                logging.warning("Unexpected image error, attempting blob fallback: %s", e)
            if image_blob is None:
                image_blob = getattr(related_part, "blob", None)
            if image_blob:
                image_blobs.append(image_blob)
        if not image_blobs:
            return None
        return image_blobs

    def __extract_table_content(self, tb):
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        def blockType(b):
            pattern = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg"),
            ]
            for p, n in pattern:
                if re.search(p, b):
                    return n
            tks = _simple_tokenize(b)
            if len(tks) > 3:
                return "Tx" if len(tks) < 12 else "Lx"
            return "Ot"

        if len(df) < 2:
            return []
        max_type = Counter(
            [blockType(str(df.iloc[i, j])) for i in range(1, len(df)) for j in range(len(df.iloc[i, :]))]
        )
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j])) for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000000):
        self.doc = Document(fnm) if isinstance(fnm, str) else Document(BytesIO(fnm))
        pn = 0
        secs = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            runs_within_single_paragraph = []
            for run in p.runs:
                if pn > to_page:
                    break
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text)
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1
            secs.append(
                ("".join(runs_within_single_paragraph), p.style.name if hasattr(p.style, "name") else "")
            )
        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        return secs, tbls
